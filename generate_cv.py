from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sys

doc = Document()

# ============ STYLES ============
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

# Heading 1 - Name
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(22)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0x53, 0xe2)  # Walmart blue
h1.paragraph_format.space_after = Pt(2)

# Heading 2 - Section titles
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0x53, 0xe2)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(6)

# Heading 3 - Job titles
h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(11)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after = Pt(2)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Add a thin line via border
    from docx.oxml.ns import qn
    from lxml import etree
    pPr = p._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0053E2')


def add_job(doc, title, company, period, description, bullets=None):
    # Job title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Company + period
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run_co = p2.add_run(company)
    run_co.italic = True
    run_co.font.size = Pt(10)
    run_co.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_sep = p2.add_run('  |  ')
    run_sep.font.size = Pt(10)
    run_sep.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run_per = p2.add_run(period)
    run_per.font.size = Pt(10)
    run_per.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Description
    if description:
        p3 = doc.add_paragraph(description)
        p3.paragraph_format.space_after = Pt(4)

    # Bullets
    if bullets:
        for bullet in bullets:
            bp = doc.add_paragraph(bullet, style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.space_before = Pt(0)


# ============ HEADER ============
p_name = doc.add_heading('Daniela Oriele Valdebenito Zarate', level=1)
p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Contact info
p_contact = doc.add_paragraph()
p_contact.paragraph_format.space_after = Pt(2)
run_email = p_contact.add_run('daniela.valdebenito.z@gmail.com')
run_email.font.size = Pt(10)
run_email.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p_contact.add_run('  |  ').font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run_phone = p_contact.add_run('+56 9 9257 0698')
run_phone.font.size = Pt(10)
run_phone.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_separator(doc)

# ============ PERFIL PROFESIONAL ============
doc.add_heading('Perfil Profesional', level=2)

perfil = (
    'Profesional de Recursos Humanos con más de 16 años de experiencia en el sector retail y logística, '
    'especializada en gestión integral de personas, desarrollo organizacional, relaciones laborales y '
    'transformación cultural. Trayectoria consolidada en Walmart Chile liderando equipos de RRHH en '
    'centros de distribución y operaciones de tiendas, con enfoque en metodologías ágiles, mejora continua '
    'y construcción de compromiso organizacional. Relacionadora Pública de formación, con diplomados en '
    'Desarrollo de Personas y Coaching Organizacional de la Universidad de Chile.'
)
doc.add_paragraph(perfil)

add_separator(doc)

# ============ EXPERIENCIA LABORAL ============
doc.add_heading('Experiencia Laboral', level=2)

# Job 1
add_job(
    doc,
    title='Jefe de Mercado Centralizado — Equipo Centralizado de Operaciones',
    company='Walmart Chile',
    period='Mayo 2022 – Actualidad',
    description=(
        'Integrante del equipo centralizado de RRHH bajo metodología Agile, con foco en facilitar '
        'la puesta en marcha de procesos y herramientas de gestión de personas para la operación en tiendas.'
    ),
    bullets=[
        'Diseño e implementación de procesos ágiles de RRHH orientados a la operación de tiendas.',
        'Desarrollo de herramientas y flujos de trabajo que optimizan la gestión diaria de personas en la operación.',
        'Articulación entre equipos corporativos de personas y la operación, asegurando alineamiento estratégico.',
        'Facilitación de la adopción de metodologías ágiles en los procesos de recursos humanos.',
    ]
)

# Job 2
add_job(
    doc,
    title='Jefe de RRHH Mercado — Super Bodega aCuenta, Operaciones',
    company='Walmart Chile',
    period='Mayo 2018 – Mayo 2022',
    description=(
        'Liderazgo del equipo de RRHH en múltiples tiendas del formato Super Bodega aCuenta, '
        'actuando como facilitador de la gestión de personas en la operación.'
    ),
    bullets=[
        'Dirección y coordinación de equipos de RRHH en tiendas, asegurando cumplimiento de procesos legales y corporativos.',
        'Entrega de herramientas de planificación semanal y mensual para la gestión de personas.',
        'Supervisión de beneficios, bienestar y cumplimiento normativo laboral.',
        'Apoyo estratégico a gerencias de tienda en estructura de dotación y gestión del talento.',
    ]
)

# Job 3
add_job(
    doc,
    title='Jefe de Área RRHH — Logística',
    company='Walmart Chile',
    period='Diciembre 2015 – Abril 2018',
    description=(
        'Responsable de la dirección y control de recursos humanos en Centros de Distribución, '
        'garantizando una administración equitativa y un servicio de calidad a clientes internos.'
    ),
    bullets=[
        'Control de procesos de reclutamiento y selección para operaciones logísticas.',
        'Dirección de programas de beneficios y bienestar en Centros de Distribución.',
        'Gestión del desempeño organizacional: seguimiento, instrumentos de evaluación y planes de mejora continua.',
        'Facilitación de relaciones con sindicatos y unidades de negocio en temas de gestión de personas.',
        'Diseño y gestión de proyectos de RRHH enfocados en eficiencias operacionales.',
        'Planificación de desarrollo profesional: capacitaciones, evaluaciones de desempeño y ascensos.',
        'Administración y control de presupuestos de RRHH (calidad de vida, alimentación, uniformes, transporte, convenios colectivos, becas).',
    ]
)

# Job 4
add_job(
    doc,
    title='Jefe Servicio a Personas — Logística',
    company='Walmart Chile',
    period='Marzo 2013 – Noviembre 2015',
    description=(
        'Dirección de procesos de RRHH en el Centro de Distribución Puerto Santiago, '
        'garantizando el cumplimiento de lineamientos corporativos y el bienestar de los colaboradores.'
    ),
    bullets=[
        'Supervisión de reclutamiento y selección conforme a descripciones de cargo definidas.',
        'Ejecución y control de la gestión del desempeño anual e identificación de brechas de competencias.',
        'Control y supervisión de beneficios corporativos (alimentación, transporte, uniformes).',
        'Soporte a operación en estructura de planta, ingresos y desvinculaciones.',
        'Liderazgo de relaciones sindicales al interior del Centro de Distribución.',
    ]
)

# Job 5
add_job(
    doc,
    title='Coordinadora de RRHH — Logística',
    company='Walmart Chile (ex D&S / Ekono)',
    period='Noviembre 2009 – Febrero 2013',
    description=(
        'Coordinación integral de procesos de RRHH en Centros de Distribución, abarcando '
        'administración de beneficios, nómina y optimización de recursos.'
    ),
    bullets=[
        'Planificación y coordinación de compra y distribución de uniformes con gerencia de compras indirectas.',
        'Gestión de información para compensaciones: descripciones de cargo y análisis de remuneraciones.',
        'Coordinación con clientes internos en apertura de vacantes, requerimientos y revisión de rentas.',
        'Control y supervisión de KPIs de nómina, beneficios y asistencia, incluyendo análisis y reportería.',
    ]
)

# Job 6
add_job(
    doc,
    title='Práctica Profesional — RRHH',
    company='D&S (actualmente Walmart Chile)',
    period='Noviembre 2008 – Marzo 2009',
    description=None,
    bullets=[
        'Apoyo en procesos de reclutamiento y selección.',
        'Colaboración en procesos de inducción laboral para nuevos colaboradores.',
    ]
)

add_separator(doc)

# ============ EDUCACIÓN ============
doc.add_heading('Educación', level=2)

p_edu = doc.add_paragraph()
run_title = p_edu.add_run('Relacionadora Pública')
run_title.bold = True
run_title.font.size = Pt(11)
p_edu2 = doc.add_paragraph()
run_uni = p_edu2.add_run('Universidad Santo Tomás')
run_uni.italic = True
run_uni.font.size = Pt(10)
run_uni.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p_edu2.add_run('  |  ').font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run_yr = p_edu2.add_run('2004 – 2008')
run_yr.font.size = Pt(10)
run_yr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_separator(doc)

# ============ FORMACIÓN COMPLEMENTARIA ============
doc.add_heading('Formación Complementaria', level=2)

formacion = [
    'Diplomado en Desarrollo de Personas y Organizaciones — Escuela de Negocios, Universidad de Chile.',
    'Diplomado en Coaching y Cambio Organizacional — Escuela de Negocios, Universidad de Chile.',
    'Gestora de Inclusión Certificada.',
    'Curso Estrategia de la Innovación — Escuela de Negocios, Universidad Adolfo Ibáñez.',
    'Seminario Metodología Hay Group: Análisis y Descripción de Cargos — Hay Group.',
    'Seminario Metodología Hay Group: Evaluación y Posicionamiento de Cargos — Hay Group.',
    'Curso Manejo de Sistema SAP — Departamento de Educación Continua, DUOC UC.',
]
for item in formacion:
    bp = doc.add_paragraph(item, style='List Bullet')
    bp.paragraph_format.space_after = Pt(2)

add_separator(doc)

# ============ CONOCIMIENTOS Y HERRAMIENTAS ============
doc.add_heading('Conocimientos y Herramientas', level=2)

conocimientos = [
    'Microsoft Excel — Nivel avanzado.',
    'SAP — Nivel intermedio.',
    'SPSS (programa estadístico) — Nivel intermedio.',
    'Inglés — Nivel intermedio.',
    'Metodologías Ágiles (Scrum/Kanban) aplicadas a RRHH.',
]
for item in conocimientos:
    bp = doc.add_paragraph(item, style='List Bullet')
    bp.paragraph_format.space_after = Pt(2)

# ============ SAVE ============
output_path = r'C:\Users\dvaldeb\OneDrive - Walmart Inc\Escritorio\CV_Daniela_Valdebenito_Profesional.docx'
doc.save(output_path)
print(f'CV guardado en: {output_path}', file=sys.stderr)
